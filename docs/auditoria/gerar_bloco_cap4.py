# -*- coding: utf-8 -*-
"""Gera o bloco pronto do Capitulo 4 (atribuicao de fatores) para colar no TCC."""
import csv
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

RAIZ = Path(__file__).resolve().parent.parent.parent
CSV = RAIZ / "data" / "Estrategias" / "apendice_I_atribuicao_fatores.csv"
SAIDA = Path(__file__).resolve().parent / "Bloco_Cap4_Atribuicao_Fatores_v2.docx"

ORDEM = [
    "EqualWeight", "EqualWeight_BuyHold", "InvVol", "MinVar", "MinVar_c10",
    "MaxSharpe", "MaxSharpe_c10", "MaxOmega", "MaxSortino", "MaxKappa3",
    "MinCVaR", "MinCDaR", "BL_classico", "BL_classico_c10",
    "BL_downside", "BL_downside_c10",
]

def br(x, casas=2):
    return f"{float(x):.{casas}f}".replace(".", ",").replace("-", "−")

dados = {}
with open(CSV, encoding="utf-8") as f:
    for linha in csv.DictReader(f):
        dados[linha["Estratégia"]] = linha

doc = Document()
estilo = doc.styles["Normal"]
estilo.font.name = "Times New Roman"
estilo.font.size = Pt(12)

def paragrafo(texto, depois=6, just=True, negrito=False, tamanho=12):
    p = doc.add_paragraph()
    if just:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(texto)
    r.font.name = "Times New Roman"
    r.font.size = Pt(tamanho)
    r.font.bold = negrito
    p.paragraph_format.space_after = Pt(depois)
    return p

paragrafo(
    "BLOCO PRONTO PARA O CAPÍTULO 4 — colar ao FINAL da Seção 4.3 (Análise "
    "Comparativa Consolidada), após o último parágrafo de análise da Tabela 8 e "
    "antes do título da Seção 4.4 (Discussão: Determinantes do Desempenho). "
    "A tabela entra como Tabela 9; a atual Tabela 9 (Síntese, Cap. 5) e as "
    "seguintes renumeram automaticamente ao atualizar os campos (Ctrl+T, F9), "
    "pois as legendas usam campo SEQ — recomenda-se recriar esta legenda via "
    "Referências > Inserir Legenda para herdar o SEQ.",
    negrito=True, just=False, tamanho=10,
)
doc.add_paragraph()

# Parágrafo 1 — metodologia e alfas
paragrafo(
    "Para decompor a origem dos retornos e isolar a real geração de valor das "
    "estratégias, estimou-se a atribuição de desempenho ex-post por meio de "
    "regressões dos excessos de retorno diários de cada carteira contra os fatores "
    "de risco brasileiros do NEFIN/USP, sob dois arranjos: o modelo de quatro "
    "fatores de Carhart (1997) — mercado (Rm−Rf), tamanho (SMB), valor (HML) e "
    "momentum (WML) — e o modelo de cinco fatores do NEFIN, que acrescenta o fator "
    "de iliquidez (IML), particularmente relevante para o mercado acionário "
    "brasileiro. Os erros-padrão são robustos a heterocedasticidade e "
    "autocorrelação (HAC de Newey-West). Conforme a Tabela 9, a carteira "
    "BL_classico apresenta o maior alfa pontual da amostra (6,12% a.a. no modelo "
    "de Carhart e 6,28% a.a. no modelo NEFIN), embora estatisticamente não "
    "significativo aos níveis convencionais (t ≈ 1,0) — resultado esperado em "
    "séries de horizonte decenal com volatilidade elevada. A única exceção é a "
    "EqualWeight_BuyHold, cujo alfa de 4,61% a.a. no modelo de cinco fatores é "
    "significativo a 5% (t = 2,02); no extremo oposto, o alfa negativo do MinCDaR "
    "(−5,35% a.a.) corrobora o diagnóstico de fragilidade da estratégia."
)

# Parágrafo 2 — cargas fatoriais
paragrafo(
    "As cargas fatoriais revelam dois padrões centrais. Primeiro, a exposição ao "
    "momentum (βWML) é substancialmente positiva apenas na família Black-Litterman "
    "(de 0,303 a 0,464), validando empiricamente que as visões construídas a "
    "partir do sinal 12-1 se traduzem, de fato, em inclinação efetiva ao fator "
    "WML — enquanto as demais estratégias exibem cargas próximas de zero. Segundo, "
    "todas as 16 carteiras apresentam exposição positiva ao prêmio de iliquidez "
    "(βIML entre 0,16 e 0,31), indicando que a otimização no mercado brasileiro "
    "tende a capturar retornos associados a ativos menos líquidos; coerentemente, "
    "a inclusão do IML eleva o R² ajustado das regressões (de 0,842 para 0,850 na "
    "EqualWeight, por exemplo), reforçando a relevância desse fator local. Em "
    "conjunto, os resultados indicam que parte do desempenho das carteiras "
    "otimizadas decorre de prêmios de risco sistemáticos — momentum e iliquidez —, "
    "e não exclusivamente de habilidade alocativa, nuance que a análise por alfa "
    "torna explícita."
)

doc.add_paragraph()

# Legenda ABNT (acima da tabela)
leg = doc.add_paragraph()
leg.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = leg.add_run(
    "Tabela 9 - Atribuição de desempenho ex-post: alfas anualizados e cargas "
    "fatoriais (modelos Carhart 4F e NEFIN 5F), 2015–2025"
)
r.font.name = "Times New Roman"
r.font.size = Pt(11)
r.font.bold = True

tab = doc.add_table(rows=1, cols=8)
tab.style = "Table Grid"
tab.alignment = WD_TABLE_ALIGNMENT.CENTER
cab = ["Estratégia", "α C4F\n(% a.a.)", "t (NW)", "βWML",
       "α N5F\n(% a.a.)", "t (NW)", "βIML", "R² aj.\n(N5F)"]
for i, txt in enumerate(cab):
    cel = tab.rows[0].cells[i]
    cel.text = ""
    p = cel.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(txt)
    r.font.name = "Times New Roman"
    r.font.size = Pt(10)
    r.font.bold = True

for est in ORDEM:
    d = dados[est]
    sig = "*" if abs(float(d["NEFIN5_t_stat"])) > 1.96 else ""
    vals = [
        est,
        br(d["Carhart_Alpha_anualized_(%)"]),
        br(d["Carhart_t_stat"]),
        br(d["Carhart_Beta_WML"], 3),
        br(d["NEFIN5_Alpha_anualized_(%)"]) + sig,
        br(d["NEFIN5_t_stat"]),
        br(d["NEFIN5_Beta_IML"], 3),
        br(d["NEFIN5_R2_adj"], 3),
    ]
    linha = tab.add_row()
    for i, v in enumerate(vals):
        cel = linha.cells[i]
        cel.text = ""
        p = cel.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i == 0 else WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(v)
        r.font.name = "Times New Roman"
        r.font.size = Pt(10)

# Fonte e nota (abaixo da tabela)
fonte = doc.add_paragraph()
fonte.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fonte.add_run(
    "Fonte: elaborada pelo autor (2026), a partir dos fatores de risco do "
    "NEFIN/USP. Nota: α = intercepto anualizado da regressão dos excessos de "
    "retorno diários; estatísticas t com erros-padrão HAC de Newey-West "
    "(5 defasagens); * significativo a 5%. C4F = Carhart 4 fatores "
    "(Rm−Rf, SMB, HML, WML); N5F = NEFIN 5 fatores (C4F + IML)."
)
r.font.name = "Times New Roman"
r.font.size = Pt(10)

doc.save(SAIDA)
print(f"OK: {SAIDA}")
