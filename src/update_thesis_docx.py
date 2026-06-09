import sys
import shutil
import pandas as pd
import numpy as np
from pathlib import Path
from docx import Document

# Set console output encoding to utf-8
sys.stdout.reconfigure(encoding='utf-8')

# Paths
project_root = Path(r"C:\VSCodeWorkspace\1_TCC_Final")
docx_orig = project_root / "docs" / "Entrega_12_Pedro_Reis_07062026_1500.docx"
docx_dest = project_root / "docs" / "Entrega_12_Pedro_Reis_07062026_1500_corrigido.docx"

# Make a backup/copy of the original
print(f"Creating corrected copy of docx at: {docx_dest}")
shutil.copy2(docx_orig, docx_dest)

# Load document
doc = Document(docx_dest)

# -----------------------------------------------------------------------------
# 1. PARAGRAPH TEXT REPLACEMENTS (RUN-LEVEL FOR FORMATTING PRESERVATION)
# -----------------------------------------------------------------------------

# We define a mapping of replacements to apply to specific paragraphs.
# Key is paragraph index, value is list of tuples (old_text, new_text) to replace.
paragraph_replacements = {
    913: [
        ("amostra final de 118 ativos", "amostra final de 102 ativos")
    ],
    946: [
        ("sobre os 118 ativos da amostra", "sobre os 102 ativos da amostra")
    ],
    947: [
        ("estacionários em todos os 118 ativos", "estacionários em todos os 102 ativos")
    ],
    948: [
        ("retornos em 118 ativos", "retornos em 102 ativos"),
        ("em **118 de 118 ativos**", "em **102 de 102 ativos**")
    ],
    949: [
        ("em **118 de 118 ativos**", "em **102 de 102 ativos**")
    ],
    950: [
        ("em **118 de 118 ativos**", "em **102 de 102 ativos**")
    ],
    1096: [
        ("sendo N=118", "sendo N=102")
    ],
    1139: [
        ("universo investível final de 118 ativos da B3", "universo investível final de 102 ativos da B3")
    ],
    1159: [
        ("obteve Sharpe de apenas 0,070, ficando aquém do próprio IBOVESPA", "obteve Sharpe de 0,328 (CAGR de 14,65% a.a.), superando o IBOVESPA"),
        ("registrou Sharpe negativo (−0,075) e CAGR de 5,97% a.a.", "registrou Sharpe de 0,290 e CAGR de 13,91% a.a."),
        ("EqualWeight_BuyHold: Sharpe 0,372, CAGR 15,71% a.a.", "EqualWeight_BuyHold: Sharpe 0,443, CAGR 17,28% a.a."),
        ("só supera o índice passivo quando mantida sem rebalanceamento", "rebalanceada supera o índice passivo (IBOVESPA Sharpe 0,178), mas perde por larga margem para sua própria versão buy-and-hold sem rebalanceamento")
    ],
    1163: [
        ("MaxSortino (Sharpe 0,263; Sortino 0,371)", "MaxSortino (Sharpe 0,281; Sortino 0,397)"),
        ("MaxKappa3 (Sharpe 0,269; Sortino 0,381)", "MaxKappa3 (Sharpe 0,284; Sortino 0,403)")
    ],
    1164: [
        ("Sharpe 0,214; volatilidade 12,96%; drawdown −26,26%", "Sharpe 0,264; volatilidade 12,88%; drawdown −26,44%")
    ],
    1165: [
        ("Sharpe de 0,111 e volatilidade anualizada de 21,21%", "Sharpe de 0,149 e volatilidade anualizada de 21,23%"),
        ("incorreu em severa degeneração algorítmica, registrando o pior desempenho global da amostra: CAGR de −1,75% a.a., Sharpe de −0,417 e drawdown máximo de −81,81%, o que traduz uma erosão de capital superior a 80%",
         "registrou o pior desempenho global da amostra, embora atenuado pelas correções de condicionamento do solver e limpeza do universo: CAGR de 5,36% a.a., Sharpe de −0,105 e drawdown máximo de −62,46%, o que traduz uma perda de capital superior a 60%")
    ],
    1169: [
        ("CAGR (21,12% a.a.)", "CAGR (24,31% a.a.)"),
        ("índice de Sharpe (0,511)", "índice de Sharpe (0,611)"),
        ("Sortino de 0,740", "Sortino de 0,890"),
        ("quase três vezes", "mais de três vezes"),
        ("CAGR de 20,66% a.a. e Sharpe de 0,456", "CAGR de 22,91% a.a. e Sharpe de 0,514")
    ],
    1170: [
        ("BL_classico: 25,99%; BL_downside: 32,36%", "BL_classico: 25,97%; BL_downside: 32,17%"),
        ("drawdowns profundos (−43,78% e −59,82%, respectivamente", "drawdowns profundos (−38,93% e −57,41%, respectivamente")
    ],
    1171: [
        ("volatilidade de 25,99% para 20,30%", "volatilidade de 25,97% para 20,20%"),
        ("drawdown de −43,78% para −34,33%", "drawdown de −38,93% para −33,60%"),
        ("0,05 ponto de Sharpe (0,511 → 0,460)", "0,05 ponto de Sharpe (0,611 → 0,563)"),
        ("CAGR (21,12% → 17,99%)", "CAGR (24,31% → 20,44%)"),
        ("volatilidade quase patológica de 32,36% para 22,32%", "volatilidade quase patológica de 32,17% para 22,30%"),
        ("drawdown de −59,82% para −37,13%", "drawdown de −57,41% para −38,01%")
    ],
    1172: [
        ("646% e 990% ao ano (BL_classico: 825%; BL_downside: 990%)", "629% e 948% ao ano (BL_classico: 798%; BL_downside: 948%)"),
        ("CAGR de 21,12% é líquido", "CAGR de 24,31% é líquido")
    ],
    1215: [
        ("CAGR entre a carteira 1/N buy-and-hold (15,71%) e a 1/N rebalanceada mensalmente (5,97%)", "CAGR entre a carteira 1/N buy-and-hold (17,28%) e a 1/N rebalanceada mensalmente (13,91%)")
    ],
    1228: [
        ("universo de 118 ações", "universo de 102 ações")
    ],
    1238: [
        ("Índice de Sharpe de 0,511", "Índice de Sharpe de 0,611"),
        ("retorno anualizado (CAGR) de 21,12%", "retorno anualizado (CAGR) de 24,31%"),
        ("próxima a 300%", "próxima a 240% em relação à métrica registrada pelo benchmark (IBOVESPA, Sharpe 0,178)")
    ],
    1239: [
        ("Sharpe de −0,417 e drawdown de −81,81% —", "Sharpe de −0,105 e drawdown de −62,46% —"),
        ("(Sharpe 0,372)", "(Sharpe 0,443)"),
        ("degenerou em destruição de valor (Sharpe −0,075)", "reduziu o retorno ajustado ao risco (Sharpe 0,290)")
    ],
    1251: [
        ("1/N rebalanceada (Sharpe −0,075) e sua versão buy-and-hold (Sharpe 0,372)", "1/N rebalanceada (Sharpe 0,290) e sua versão buy-and-hold (Sharpe 0,443)")
    ],
    1401: [
        ("N = 118", "N = 102")
    ],
    1403: [
        ("Relação dos 118 ativos que", "Relação dos 102 ativos que")
    ],
    1416: [
        ("reteve 118 ativos que atenderam a este critério de alta liquidez.", "reteve 118 ativos no critério de alta liquidez (ADTV). Adicionalmente, na Etapa VII, o Classificador de Integridade Point-in-Time avaliou a consistência das séries históricas através de reorganizações societárias, penny stocks persistentes e situações especiais (CODAHIST), excluindo 16 ativos adicionais e estabelecendo o universo final de 102 ativos.")
    ]
}

# Apply paragraph replacements by iterating over runs
print("Applying text replacements to paragraphs and preserving formatting...")
for idx, replacements in paragraph_replacements.items():
    para = doc.paragraphs[idx]
    for old, new in replacements:
        replaced = False
        # Try replacing in run texts first
        for run in para.runs:
            if old in run.text:
                run.text = run.text.replace(old, new)
                replaced = True
        
        # If it wasn't replaced (perhaps because old text is split across multiple runs),
        # we do a fallback paragraph-level replacement
        if not replaced and old in para.text:
            print(f"  [Fallback] Replacing in Para {idx} text directly: '{old}' -> '{new}'")
            # Clear runs except the first, put all text in first run
            # Note: This is a backup fallback to ensure replacement occurs.
            para.text = para.text.replace(old, new)

# -----------------------------------------------------------------------------
# 2. TABLES 7 AND 8 (TABELAS 8 E 9): PERFORMANCE METRICS
# -----------------------------------------------------------------------------

# Load new performance metrics from apendice_H_painel_metricas.csv and desempenho_estrategias.csv
desempenho_path = project_root / "data" / "Estrategias" / "desempenho_estrategias.csv"
des_df = pd.read_csv(desempenho_path).rename(columns={"Unnamed: 0": "Estrategia"})

# Helper to format numbers into Portuguese style (comma as decimal separator)
def fmt_pt(val, decimals=2, is_percentage=False, is_sharpe_sortino=False):
    if pd.isna(val):
        return "—"
    if is_percentage:
        val = val * 100
    formatted = f"{val:.{decimals}f}".replace('.', ',')
    return formatted

# Let's map strategy names from CSV to docx Table 7 and Table 8
# We create a dictionary of new metrics for each strategy
metrics_dict = {}
for _, row in des_df.iterrows():
    est = row["Estrategia"]
    # For IBOVESPA, we map IBOV
    key = "IBOVESPA" if est == "IBOV" else est
    # CAGR and Vol are in ratio in CSV, we need percentage
    # Sharpe, Sortino, MaxDD, Turnover
    metrics_dict[key] = {
        "cagr": fmt_pt(row["ret_anual"], decimals=2, is_percentage=True),
        "vol": fmt_pt(row["vol_anual"], decimals=2, is_percentage=True),
        "sharpe": fmt_pt(row["sharpe"], decimals=3),
        "sortino": fmt_pt(row["sortino"], decimals=3),
        "mdd": fmt_pt(row["max_drawdown"], decimals=2, is_percentage=True),
        "turnover": fmt_pt(row["turnover_aa"], decimals=1, is_percentage=True)
    }

# Also need Accum. Return (Ret. Acum. %) which is in apendice_H_painel_metricas.csv
painel_path = project_root / "data" / "Estrategias" / "apendice_H_painel_metricas.csv"
painel_df = pd.read_csv(painel_path)
for _, row in painel_df.iterrows():
    est = row["Estratégia"]
    if est in metrics_dict:
        metrics_dict[est]["ret_acum"] = fmt_pt(row["Retorno Acumulado (%)"] / 100, decimals=1, is_percentage=True)

# Add IBOVESPA Accum. Return if not present
metrics_dict["IBOVESPA"]["ret_acum"] = fmt_pt(painel_df[painel_df["Estratégia"] == "IBOVESPA"]["Retorno Acumulado (%)"].values[0] / 100, decimals=1, is_percentage=True)

# Let's process Table 7 (Tabela 8 in the text)
# Row 0 is header. Rows 1 to 17 are the 17 strategies.
# We will read their strategy names, fetch new metrics, update them, sort them by Sharpe descending,
# and write them back in the sorted order.
table_7 = doc.tables[7]
print("\nUpdating Table 7 (Tabela 8: Out-of-Sample Performance)...")

table_7_rows_data = []
for r_idx in range(1, len(table_7.rows)):
    row = table_7.rows[r_idx]
    strat_name = row.cells[0].text.strip()
    fam = row.cells[1].text.strip()
    
    # Get metrics
    if strat_name in metrics_dict:
        m = metrics_dict[strat_name]
        # Keep numeric Sharpe for sorting
        sharpe_num = des_df[des_df["Estrategia"] == ("IBOV" if strat_name == "IBOVESPA" else strat_name)]["sharpe"].values[0]
        
        row_data = {
            "name": strat_name,
            "fam": fam,
            "ret_acum": m["ret_acum"],
            "cagr": m["cagr"],
            "vol": m["vol"],
            "sharpe": m["sharpe"],
            "sortino": m["sortino"],
            "mdd": m["mdd"],
            "turnover": m["turnover"],
            "sharpe_num": sharpe_num
        }
        table_7_rows_data.append(row_data)
    else:
        print(f"  [Warning] Strategy {strat_name} not found in metrics!")

# Sort Table 7 rows by Sharpe descending
table_7_rows_data.sort(key=lambda x: x["sharpe_num"], reverse=True)

# Write back in sorted order
for idx, data in enumerate(table_7_rows_data):
    row_idx = idx + 1
    row = table_7.rows[row_idx]
    row.cells[0].text = data["name"]
    row.cells[1].text = data["fam"]
    row.cells[2].text = data["ret_acum"]
    row.cells[3].text = data["cagr"]
    row.cells[4].text = data["vol"]
    row.cells[5].text = data["sharpe"]
    row.cells[6].text = data["sortino"]
    row.cells[7].text = data["mdd"]
    row.cells[8].text = data["turnover"]
    print(f"  Row {row_idx}: {data['name']} (Sharpe: {data['sharpe']})")


# Let's process Table 8 (Tabela 9 in the text)
# Row 0 is header. Rows 1 to 9 are the 9 comparison strategies.
table_8 = doc.tables[8]
print("\nUpdating Table 8 (Tabela 9: Key Strategies Comparison)...")

# Map display names in Table 8 to keys in metrics_dict
table_8_mapping = {
    "BL_classico (BL, momentum)": "BL_classico",
    "BL_classico_c10": "BL_classico_c10",
    "BL_downside": "BL_downside",
    "EqualWeight_BuyHold (1/N comprar-e-manter)": "EqualWeight_BuyHold",
    "MaxSharpe (máxima utilidade)": "MaxSharpe",
    "MinVar (mínima variância)": "MinVar",
    "IBOVESPA (benchmark)": "IBOVESPA",
    "EqualWeight (1/N rebalanceada)": "EqualWeight",
    "MinCDaR": "MinCDaR"
}

table_8_rows_data = []
for r_idx in range(1, len(table_8.rows)):
    row = table_8.rows[r_idx]
    display_name = row.cells[0].text.strip()
    
    if display_name in table_8_mapping:
        strat_key = table_8_mapping[display_name]
        m = metrics_dict[strat_key]
        sharpe_num = des_df[des_df["Estrategia"] == ("IBOV" if strat_key == "IBOVESPA" else strat_key)]["sharpe"].values[0]
        
        row_data = {
            "display_name": display_name,
            "cagr": m["cagr"],
            "vol": m["vol"],
            "sharpe": m["sharpe"],
            "mdd": m["mdd"],
            "turnover": m["turnover"],
            "sharpe_num": sharpe_num
        }
        table_8_rows_data.append(row_data)
    else:
        print(f"  [Warning] Table 8 display name '{display_name}' not mapped!")

# Sort Table 8 rows by Sharpe descending
table_8_rows_data.sort(key=lambda x: x["sharpe_num"], reverse=True)

# Write back in sorted order
for idx, data in enumerate(table_8_rows_data):
    row_idx = idx + 1
    row = table_8.rows[row_idx]
    row.cells[0].text = data["display_name"]
    row.cells[1].text = data["cagr"]
    row.cells[2].text = data["vol"]
    row.cells[3].text = data["sharpe"]
    row.cells[4].text = data["mdd"]
    row.cells[5].text = data["turnover"]
    print(f"  Row {row_idx}: {data['display_name']} (Sharpe: {data['sharpe']})")


# -----------------------------------------------------------------------------
# 3. TABLE 10 (TABELA 11: SECTOR COMPOSITION)
# -----------------------------------------------------------------------------
print("\nUpdating Table 10 (Tabela 11: Sector Composition)...")
table_10 = doc.tables[10]

new_sectors = [
    ("Consumo Cíclico", 20, "19,6%"),
    ("Utilidade Pública", 16, "15,7%"),
    ("Materiais Básicos", 16, "15,7%"),
    ("Financeiro", 14, "13,7%"),
    ("Bens Industriais", 11, "10,8%"),
    ("Consumo Não Cíclico", 10, "9,8%"),
    ("Saúde", 5, "4,9%"),
    ("Comunicações", 4, "3,9%"),
    ("Petróleo, Gás e Biocombustíveis", 3, "2,9%"),
    ("Tecnologia da Informação", 3, "2,9%")
]

for idx, (sector, count, pct) in enumerate(new_sectors):
    row_idx = idx + 1
    row = table_10.rows[row_idx]
    row.cells[0].text = sector
    row.cells[1].text = str(count)
    row.cells[2].text = pct
    print(f"  Row {row_idx}: {sector} | {count} | {pct}")

# Update total row
total_row = table_10.rows[11]
total_row.cells[0].text = "Total"
total_row.cells[1].text = "102"
total_row.cells[2].text = "100,0%"
print(f"  Row 11: Total | 102 | 100,0%")


# -----------------------------------------------------------------------------
# 4. TABLE 11 (TABELA 12: APPENDIX A ASSET RELATION)
# -----------------------------------------------------------------------------
print("\nUpdating Table 11 (Appendix A: Asset Relation) - Removing 16 Excluded Tickers...")
table_11 = doc.tables[11]

excluded_tickers = {
    "AMER3", "ETER3", "FICT3", "GOLL54", "LIGT3", "LUPA3", "NEXP3", "OIBR3",
    "OIBR4", "PDGR3", "PDTC3", "PMAM3", "RPMG3", "RSID3", "VIVR3", "VSTE3"
}

# Find rows to remove
rows_to_remove = []
for idx, row in enumerate(table_11.rows):
    if idx == 0:
        continue
    ticker = row.cells[1].text.strip()
    if ticker in excluded_tickers:
        print(f"  Found excluded ticker {ticker} at row {idx}")
        rows_to_remove.append(row)

# Delete rows
for row in rows_to_remove:
    table_11._tbl.remove(row._tr)

print(f"  Removed {len(rows_to_remove)} rows. Remaining rows: {len(table_11.rows)}")

# Re-number the "Posição" column
for idx, row in enumerate(table_11.rows):
    if idx == 0:
        continue
    row.cells[0].text = str(idx)

# Save the updated document
doc.save(docx_dest)
print(f"\nSuccessfully saved updated document to: {docx_dest}")
