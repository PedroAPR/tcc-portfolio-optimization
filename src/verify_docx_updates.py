import sys
from pathlib import Path
from docx import Document

# Set console output encoding to utf-8
sys.stdout.reconfigure(encoding='utf-8')

project_root = Path(r"C:\VSCodeWorkspace\1_TCC_Final")
docx_path = project_root / "docs" / "Entrega_12_Pedro_Reis_07062026_1500_corrigido.docx"

if not docx_path.exists():
    print(f"Error: {docx_path} does not exist!")
    sys.exit(1)

doc = Document(docx_path)
failed = False

# 1. Check Paragraphs for old text
print("=== VERIFYING PARAGRAPHS ===")
check_paras = {
    913: "102",
    946: "102",
    947: "102",
    948: "102 de 102",
    949: "102 de 102",
    950: "102 de 102",
    1096: "N=102",
    1139: "102",
    1159: "0,290",
    1163: "0,281",
    1164: "0,264",
    1165: "−0,105",
    1169: "24,31% a.a.",
    1170: "25,97%",
    1171: "0,611",
    1172: "24,31%",
    1215: "13,91%",
    1228: "102 ações",
    1238: "0,611",
    1239: "−0,105",
    1251: "0,290",
    1401: "N = 102",
    1403: "102",
    1416: "102"
}

for idx, expected in check_paras.items():
    text = doc.paragraphs[idx].text
    if expected not in text:
        print(f"  [FAIL] Para {idx} does not contain '{expected}'!")
        print(f"  Actual text: {text[:150]}...")
        failed = True
    else:
        print(f"  [OK] Para {idx} verified with '{expected}'")

# 2. Check Table 7 (Tabela 8: Performance)
print("\n=== VERIFYING TABLE 7 (TABELA 8) ===")
table_7 = doc.tables[7]
row_1 = [c.text.strip() for c in table_7.rows[1].cells]
row_17 = [c.text.strip() for c in table_7.rows[17].cells]

if row_1[0] != "BL_classico" or row_1[5] != "0,611":
    print(f"  [FAIL] Table 7 row 1 is wrong: {row_1}")
    failed = True
else:
    print(f"  [OK] Table 7 row 1 verified: {row_1}")

if row_17[0] != "MinCDaR" or row_17[5] != "-0,105":
    print(f"  [FAIL] Table 7 row 17 is wrong: {row_17}")
    failed = True
else:
    print(f"  [OK] Table 7 row 17 verified: {row_17}")

# 3. Check Table 8 (Tabela 9: Comparison)
print("\n=== VERIFYING TABLE 8 (TABELA 9) ===")
table_8 = doc.tables[8]
row_1_t8 = [c.text.strip() for c in table_8.rows[1].cells]
row_9_t8 = [c.text.strip() for c in table_8.rows[9].cells]

if row_1_t8[0] != "BL_classico (BL, momentum)" or row_1_t8[3] != "0,611":
    print(f"  [FAIL] Table 8 row 1 is wrong: {row_1_t8}")
    failed = True
else:
    print(f"  [OK] Table 8 row 1 verified: {row_1_t8}")

if row_9_t8[0] != "MinCDaR" or row_9_t8[3] != "-0,105":
    print(f"  [FAIL] Table 8 row 9 is wrong: {row_9_t8}")
    failed = True
else:
    print(f"  [OK] Table 8 row 9 verified: {row_9_t8}")

# 4. Check Table 10 (Tabela 11: Sector)
print("\n=== VERIFYING TABLE 10 (TABELA 11) ===")
table_10 = doc.tables[10]
total_row = [c.text.strip() for c in table_10.rows[11].cells]
if total_row[1] != "102" or total_row[2] != "100,0%":
    print(f"  [FAIL] Table 10 total row is wrong: {total_row}")
    failed = True
else:
    print(f"  [OK] Table 10 total row verified: {total_row}")

# 5. Check Table 11 (Appendix A)
print("\n=== VERIFYING TABLE 11 (APPENDIX A) ===")
table_11 = doc.tables[11]
print(f"  Total rows: {len(table_11.rows)}")
last_row = [c.text.strip() for c in table_11.rows[-1].cells]
if len(table_11.rows) != 103 or last_row[0] != "102" or last_row[1] != "YDUQ3":
    print(f"  [FAIL] Table 11 is wrong! Rows={len(table_11.rows)}, Last row={last_row}")
    failed = True
else:
    print(f"  [OK] Table 11 verified. Rows={len(table_11.rows)}, Last row={last_row}")

if failed:
    print("\n>>> VERIFICATION FAILED! <<<")
    sys.exit(1)
else:
    print("\n>>> ALL VERIFICATIONS PASSED SUCCESSFULLY! <<<")
    sys.exit(0)
