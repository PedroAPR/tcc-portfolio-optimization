import json
import sys
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# Set console output encoding to utf-8
sys.stdout.reconfigure(encoding='utf-8')

# Paths
project_root = Path(r"C:\VSCodeWorkspace\1_TCC_Final")
src_dir = project_root / "src"
docx_path = project_root / "docs" / "Apendice_Codigo_Fonte_Pipeline.docx"

print(f"Generating single Word document at: {docx_path}")

doc = Document()

# Define standard styles
styles = doc.styles
normal_style = styles['Normal']
normal_style.font.name = 'Times New Roman'
normal_style.font.size = Pt(12)
normal_style.font.color.rgb = RGBColor(0, 0, 0)
normal_style.paragraph_format.line_spacing = 1.5
normal_style.paragraph_format.space_after = Pt(6)

# Title Page Section
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run("APÊNDICE COMPUTACIONAL\n\nCÓDIGO FONTE E DOCUMENTAÇÃO DO PIPELINE DE OTIMIZAÇÃO DE CARTEIRAS")
title_run.font.name = 'Times New Roman'
title_run.font.size = Pt(16)
title_run.bold = True
title_p.paragraph_format.space_before = Pt(36)
title_p.paragraph_format.space_after = Pt(24)

intro_p = doc.add_paragraph(
    "Este apêndice reúne de forma sequencial a documentação (Markdown) e o código-fonte (Python) "
    "de todas as etapas do pipeline quantitativo de otimização de carteiras. O fluxo de processamento "
    "abrange desde a ingestão dos dados brutos das ações da B3, aplicação de filtros de liquidez e "
    "integridade, alinhamento de datas, winsorização de outliers, estimação de matrizes de covariância "
    "Ledoit-Wolf, simulações de fronteira eficiente, otimizações out-of-sample MPT/PMPT e testes "
    "econométricos via bootstrap estacionário."
)

# -----------------------------------------------------------------------------
# PIPELINE STRUCTURE TABLE
# -----------------------------------------------------------------------------
doc.add_heading("Estrutura Geral de Execução do Pipeline", level=1)
doc.add_paragraph(
    "A tabela abaixo descreve a ordem cronológica e a finalidade de cada notebook e script "
    "orquestrado para a geração dos resultados empíricos da pesquisa:"
)

stages_table_data = [
    ("1", "Conversão Parquet", "Notebook", "01_01_convertendo_em_parquet_v3.ipynb", "Ingestão e padronização binária de planilhas B3"),
    ("2", "Consolidação Dados", "Notebook", "02_01_consolidando_dados.ipynb", "Estruturação em painel unificado MultiIndex"),
    ("3a", "Pré-Integridade", "Notebook", "03_01a_Pre_Integridade.ipynb", "Filtros de presença, IPO e liquidez ADTV (p10)"),
    ("3b", "Classificador", "Notebook", "Apendice_Classificador_Integridade_Universo_v2.ipynb", "Auditoria de penny stocks e distress (COTAHIST)"),
    ("3c", "Pós-Integridade", "Notebook", "03_01c_Pos_Integridade.ipynb", "Aplicação das 16 exclusões societárias de integridade"),
    ("3d", "Orquestrador Local", "Script", "run_etapa03.py", "Sincronização local das três sub-etapas de integridade"),
    ("4", "Taxas Livres Risco", "Notebook", "04_01_Taxas_Livres_Risco_SGS_Final.ipynb", "Ingestão do CDI e Selic diários via SGS/Banco Central"),
    ("5", "Alinhamento Retornos", "Notebook", "05_01_Alinhamento_e_Retornos.ipynb", "Alinhamento cronológico das ações com IBOV/CDI"),
    ("6", "Saneamento Winsor", "Notebook", "05_02_Saneamento_e_Winsorizacao.ipynb", "Winsorização de cauda e saneamento de outliers (MAD)"),
    ("7", "Covariância Ledoit-Wolf", "Notebook", "06_01_Estimacao_LedoitWolf.ipynb", "Estimação e encolhimento de covariância para 102 ativos"),
    ("8", "Backtest OOS", "Notebook", "07_01_Otimizacao_Carteiras.ipynb", "Simulação das 16 carteiras ótimas com custos (50 bps)"),
    ("9", "Fronteira Eficiente", "Notebook", "08_01_Fronteira_Eficiente.ipynb", "Simulação Monte Carlo e curvas MPT/PMPT in-sample"),
    ("10", "Inferência Estatística", "Notebook", "09_01_Inferencia_Econometrica.ipynb", "Testes de Sharpe/Sortino via bootstrap e CAPM HAC")
]

table = doc.add_table(rows=1, cols=5)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Ordem'
hdr_cells[1].text = 'Estágio'
hdr_cells[2].text = 'Tipo'
hdr_cells[3].text = 'Arquivo'
hdr_cells[4].text = 'Descrição'

# Header styling
for cell in hdr_cells:
    cell.paragraphs[0].runs[0].font.bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(10)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="E0E0E0"/>')
    cell._tc.get_or_add_tcPr().append(shd)

for order, name, t_type, filename, desc in stages_table_data:
    row_cells = table.add_row().cells
    row_cells[0].text = order
    row_cells[1].text = name
    row_cells[2].text = t_type
    row_cells[3].text = filename
    row_cells[4].text = desc
    
    # Cell formatting
    for cell in row_cells:
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(9.5)
        # Add light border XML manually or let Word default handle it

# Page Break
doc.add_page_break()

# -----------------------------------------------------------------------------
# STAGE CHRONOLOGICAL INTEGRATION
# -----------------------------------------------------------------------------

pipeline_files = [
    {"type": "notebook", "path": src_dir / "01_Conversao_Parquet" / "01_01_convertendo_em_parquet_v3.ipynb", "title": "Estágio 1 — Conversão Parquet de Arquivos Brutos"},
    {"type": "notebook", "path": src_dir / "02_Consolidacao_Dados" / "02_01_consolidando_dados.ipynb", "title": "Estágio 2 — Consolidação Temporal em Painel MultiIndex"},
    {"type": "notebook", "path": src_dir / "03_Filtro_Liquidez" / "03_01a_Pre_Integridade.ipynb", "title": "Estágio 3a — Ingestão e Filtros de Liquidez ADTV"},
    {"type": "notebook", "path": src_dir / "11_Classificador_Integridade" / "Apendice_Classificador_Integridade_Universo_v2.ipynb", "title": "Estágio 3b — Auditoria Point-in-Time de Integridade de Ativos"},
    {"type": "notebook", "path": src_dir / "03_Filtro_Liquidez" / "03_01c_Pos_Integridade.ipynb", "title": "Estágio 3c — Sanitização e Aplicação da Lista de Excluídos"},
    {"type": "script", "path": src_dir / "run_etapa03.py", "title": "Estágio 3d — Script Orquestrador da Etapa de Integridade"},
    {"type": "notebook", "path": src_dir / "04_Taxas_Livres_Risco" / "04_01_Taxas_Livres_Risco_SGS_Final.ipynb", "title": "Estágio 4 — Ingestão e Alinhamento das Taxas Selic e CDI"},
    {"type": "notebook", "path": src_dir / "05_Alinhamento_Winsorizacao" / "05_01_Alinhamento_e_Retornos.ipynb", "title": "Estágio 5.1 — Alinhamento Cronológico e Cálculo de Retornos"},
    {"type": "notebook", "path": src_dir / "05_Alinhamento_Winsorizacao" / "05_02_Saneamento_e_Winsorizacao.ipynb", "title": "Estágio 5.2 — Saneamento Estatístico de Retornos e Winsorização Robustos"},
    {"type": "notebook", "path": src_dir / "06_Estimacao_Covariancia" / "06_01_Estimacao_LedoitWolf.ipynb", "title": "Estágio 6 — Estimação de Momentos Estatísticos e Ledoit-Wolf"},
    {"type": "notebook", "path": src_dir / "07_Otimizacao_Carteiras" / "07_01_Otimizacao_Carteiras.ipynb", "title": "Estágio 7 — Otimização de Carteiras e Simulação de Backtest OOS"},
    {"type": "notebook", "path": src_dir / "08_Fronteira_Eficiente" / "08_01_Fronteira_Eficiente.ipynb", "title": "Estágio 8 — Fronteiras Eficientes e Simulação Monte Carlo"},
    {"type": "notebook", "path": src_dir / "09_Inferencia_Econometrica" / "09_01_Inferencia_Econometrica.ipynb", "title": "Estágio 9 — Inferência Econométrica e Robustez Estatística"},
]

def add_code_block(doc_obj, code_text):
    """Adds a beautiful code block using a single-cell table with a grey background."""
    tbl = doc_obj.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    
    # Set grey background color (#F7F7F7)
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F7F7F7"/>')
    cell._tc.get_or_add_tcPr().append(shd)
    
    # Remove borders (optional, or thin grey border)
    borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="D3D3D3"/>'
        f'</w:tcBorders>'
    )
    cell._tc.get_or_add_tcPr().append(borders)
    
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.05
    
    run = p.add_run(code_text)
    run.font.name = 'Courier New'
    run.font.size = Pt(8.5)

def format_markdown_text(doc_obj, text):
    """Processes markdown block text and adds formatted paragraphs to the document."""
    lines = text.strip().split("\n")
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
        
        # Check for headings
        if line_strip.startswith("# "):
            h = doc_obj.add_paragraph()
            h.paragraph_format.space_before = Pt(12)
            h.paragraph_format.space_after = Pt(6)
            run = h.add_run(line_strip[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(14)
            run.bold = True
        elif line_strip.startswith("## "):
            h = doc_obj.add_paragraph()
            h.paragraph_format.space_before = Pt(10)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(line_strip[3:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(13)
            run.bold = True
        elif line_strip.startswith("### "):
            h = doc_obj.add_paragraph()
            h.paragraph_format.space_before = Pt(8)
            h.paragraph_format.space_after = Pt(4)
            run = h.add_run(line_strip[4:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)
            run.bold = True
        # Check for lists
        elif line_strip.startswith("* ") or line_strip.startswith("- "):
            p = doc_obj.add_paragraph(style='List Bullet')
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line_strip[2:])
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11.5)
        else:
            # Normal paragraph
            p = doc_obj.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            p.paragraph_format.line_spacing = 1.25  # Slightly tighter for code appendix description
            run = p.add_run(line)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)

for idx, stage in enumerate(pipeline_files, 1):
    file_type = stage["type"]
    file_path = stage["path"]
    title = stage["title"]
    
    print(f"Processing stage {idx}/{len(pipeline_files)}: {title} ({file_path.name})")
    
    # Add Section Heading
    h = doc.add_paragraph()
    h.paragraph_format.space_before = Pt(24)
    h.paragraph_format.space_after = Pt(8)
    run = h.add_run(f"Apêndice A.{idx} — {title}")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14)
    run.bold = True
    
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(12)
    p_run = p.add_run(f"Caminho do arquivo: src/{file_path.relative_to(src_dir)}")
    p_run.font.italic = True
    p_run.font.size = Pt(10)
    
    if not file_path.exists():
        print(f"  [-] File not found: {file_path.name}")
        continue
        
    if file_type == "notebook":
        with open(file_path, "r", encoding="utf-8") as f:
            nb_data = json.load(f)
            
        for c_idx, cell in enumerate(nb_data["cells"]):
            cell_source = "".join(cell["source"])
            if not cell_source.strip():
                continue
                
            if cell["cell_type"] == "markdown":
                format_markdown_text(doc, cell_source)
            elif cell["cell_type"] == "code":
                # Add code block description (optional or standard box)
                add_code_block(doc, cell_source)
    else:
        # Script (raw python script)
        with open(file_path, "r", encoding="utf-8") as f:
            script_content = f.read()
        
        # Add the entire script as a single code block
        doc.add_paragraph("Este é o código de orquestração local que sequencia os notebooks da etapa:")
        add_code_block(doc, script_content)
        
    # Page Break between sections
    doc.add_page_break()

# Save document
doc.save(docx_path)
print(f"\nDocument generation completed successfully! Saved as {docx_path}")
